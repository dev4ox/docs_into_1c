from pathlib import Path
from datetime import datetime
from llama_cpp import Llama
import settings
import os
import json
import re
import collections
import pandas as pd
import docx
import subprocess
from pdf2image import convert_from_path
import pdfplumber

import pytesseract
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

input_prompt = '''
Задача – анализ текста и извлечение параметров.
Выводи найденные товары, имеющие характеристики в формате JSON строго по инструкции. Обычно товары имеют следующие характеристики:
"Номенклатура", "Мощность, Вт", "Св. поток, Лм", "IP", "Длина, мм", "Ширина, мм", "Высота, мм", "Габариты", "Рассеиватель", "Цвет. температура, К", "Вес, кг", "Напряжение, В", "Срок службы (работы) светильника", "Температура эксплуатации", "Материал корпуса", "Тип", "Тип КСС", "Род тока", "Гарантия", "Индекс цветопередачи (CRI, Ra)", "Класс защиты от поражения электрическим током", "Коэффициент мощности (Pf)", "С регулятором яркости (диммирование)", "Ударопрочность", "Класс взрывозащиты (Ex)", "Класс пожароопасности", "Цвет корпуса", "Коэффициент пульсаций", "Прочее".
Может быть более одного товара на странице, выводи последовательно товары в формате JSON.

Если значение выражено диапазоном или с квалификаторами (например, "не более", "не менее", "от X до Y", "±10", "+-10", "около"), включай всю фразу с единицами измерения.
Если параметр отсутствует или его значение не может быть корректно извлечено, верни "не указано".
Если есть дополнительные характеристики товара, не подходящие под обычные характеристики, помести их в характеристику "Прочее".

Пример 1:
Входной текст:
"Наименование продукции: Прожектор светодиодный ASD СДО-2-20 20W или аналог. Энергопотребление, не более, Вт: 20; Входное напряжение: 85-265 В; Цветовая температура, К, не менее: 6500; Коэффициент пульсаций, не более: 5%; Угол свечения: 120°; Степень защиты, не менее IP: 65; Световой поток, Лм: не менее 1600; Габаритные размеры (L, b, h): 178*100*138; Время работы, не менее: 50 000 часов; Кронштейн крепления."
Вывод:
{
  "Номенклатура": "Прожектор светодиодный ASD СДО-2-20 20W или аналог",
  "Мощность, Вт": "не более 20 Вт",
  "Св. поток, Лм": "не менее 1600 Лм",
  "IP": "не менее 65",
  "Длина, мм": "178",
  "Ширина, мм": "100",
  "Высота, мм": "138",
  "Габариты": "178*100*138",
  "Рассеиватель": "не указано",
  "Цвет. температура, К": "не менее 6500",
  "Вес, кг": "не указано",
  "Напряжение, В": "85-265 В",
  "Срок службы (работы) светильника": "не менее 50 000 часов",
  "Температура эксплуатации": "не указано",
  "Материал корпуса": "не указано",
  "Тип": "не указано",
  "Тип КСС": "120°",
  "Род тока": "не указано",
  "Гарантия": "не указано",
  "Индекс цветопередачи (CRI, Ra)": "не указано",
  "Класс защиты от поражения электрическим током": "не указано",
  "Коэффициент мощности (Pf)": "не указано",
  "С регулятором яркости (диммирование)": "не указано",
  "Ударопрочность": "не указано",
  "Класс взрывозащиты (Ex)": "не указано",
  "Класс пожароопасности": "не указано",
  "Цвет корпуса": "не указано",
  "Коэффициент пульсаций": "не указано",
  "Прочее": "Кронштейн крепления"
}
    '''


# Маленькая для OCR, 3Gb vram, с парсером работает збс!
def extract_gemma_2_2b_it_IQ3_M(text, final_columns):
    """
    Функция обрабатывает текст с помощью LLM модели Gemma 2, формирует корректный промпт,
    отправляет запрос и извлекает JSON-ответ.
    """
    llm = Llama(
        model_path="/models/gemma/gemma-2-2b-it-IQ3_M.gguf",
        n_ctx=8192,
        n_gpu_layers=-1,
        verbose=True,
    )
    prompt = f"<start_of_turn>user\n{input_prompt}\n\nText:\n{text}\n\nJSON:<end_of_turn>\n<start_of_turn>model\n"
    output = llm(
        prompt=prompt,
        max_tokens=2048,
        temperature=0.0,
        stop=["<end_of_turn>"]  # Останавливаем генерацию после ответа
    )

    # Извлекаем текст и определяем только JSON-объект с помощью регулярного выражения
    result_text = output["choices"][0]["text"].strip()
    match = re.search(r'\{.*\}', result_text, re.DOTALL)
    if match:
        result_text = match.group(0)

    try:
        data = json.loads(result_text)
    except Exception as e:
        print("Error parsing JSON:", e)
        print("Raw downloads:", result_text)
        data = {col: "не указано" for col in final_columns}
    return data

# Не увидел отличий
def extract_gemma_2_9b_it_Q4_K_M(text, final_columns):
    llm = Llama(
        model_path="/models/gemma/gemma-2-9b-it-Q4_K_M.gguf",
        n_ctx=8192,
        n_gpu_layers=-1,
        verbose=True,
    )
    prompt = input_prompt + "\n\nText:\n" + text + "\n\nJSON:"
    output = llm(prompt=prompt, max_tokens=4096, temperature=0.0)
    result_text = output["choices"][0]["text"].strip()
    # Извлекаем только JSON-объект с помощью регулярного выражения
    match = re.search(r'\{.*\}', result_text, re.DOTALL)
    if match:
        result_text = match.group(0)
    try:
        data = json.loads(result_text)
    except Exception as e:
        print("Error parsing JSON:", e)
        print("Raw downloads:", result_text)
        data = {col: "не указано" for col in final_columns}
    return data


# Слишком большая
def extract_with_mistral(text, final_columns):
    llm = Llama(
        model_path="/models/mistral/Mistral-Nemo-Instruct-2407-Q4_K_M.gguf",
        n_ctx=4096,
        n_gpu_layers=-1,
    )
    prompt = input_prompt + "\n\nДанные для обработки:\n" + text
    output = llm(prompt=prompt, max_tokens=2048, temperature=0.1)
    result_text = output["choices"][0]["text"].strip()
    try:
        data = json.loads(result_text)
    except Exception as e:
        print("Error parsing JSON:", e)
        print("Raw downloads:", result_text)
        data = {col: "не указано" for col in final_columns}
    return data


def generate_filename(prefix: str="Форма2", ext: str=".xlsx"):
    timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
    return f"{prefix}-{timestamp}{ext}"


# Используем ExcelWriter в режиме добавления (append)
def append_df_to_excel(filename, df, sheet_name='Sheet1'):
    if os.path.exists(filename):
        with pd.ExcelWriter(filename, engine='openpyxl', mode='a', if_sheet_exists='overlay') as writer:
            startrow = writer.sheets[sheet_name].max_row
            df.to_excel(writer, sheet_name=sheet_name, index=False, header=False, startrow=startrow)
    else:
        df.to_excel(filename, index=False, sheet_name=sheet_name)


class UnifiedExcelParser:
    PRODUCT_NAMES = settings.product_names

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []

    def is_product_name(self, text):
        # Проверяет, начинается ли текст со слова из PRODUCT_NAMES
        return any(text.lower().startswith(name.lower()) for name in self.PRODUCT_NAMES)

    def contains_product_name(self, text):
        # Обязательная проверка: текст должен содержать хотя бы одно слово из PRODUCT_NAMES
        return any(name.lower() in text.lower() for name in self.PRODUCT_NAMES)

    def detect_engine(self):
        ext = self.file_path.suffix.lower()
        if ext in ['.xlsx', '.xlsm']:
            return 'openpyxl'
        elif ext == '.xls':
            return 'xlrd'
        else:
            return 'openpyxl'

    def parse_excel(self):
        if not self.file_path.exists():
            print(f"File not found: {self.file_path}")
            return
        engine = self.detect_engine()
        df = pd.read_excel(self.file_path, header=None, engine=engine)
        if df.shape[1] == 1:
            self.parse_single_column(df)
        else:
            name_column = None
            for col in df.columns:
                for row_idx in range(min(15, len(df))):
                    cell_value = str(df.iloc[row_idx, col]).lower() if pd.notna(df.iloc[row_idx, col]) else ""
                    if "наименование" in cell_value:
                        name_column = col
                        break
                if name_column is not None:
                    break
            if name_column is None:
                self.parse_single_column(df)
            else:
                extra_char = (name_column + 2 < df.shape[1])
                self.parse_multi_column(df, name_column, extra_char)

    def parse_single_column(self, df):
        current_text = ""
        for index, row in df.iterrows():
            cell_value = str(row[0]).strip()
            if self.is_product_name(cell_value):
                if current_text and self.contains_product_name(current_text):
                    self.data.append({"text": current_text})
                current_text = cell_value
            else:
                current_text += " " + cell_value
        if current_text and self.contains_product_name(current_text):
            self.data.append({"text": current_text})

    def parse_multi_column(self, df, name_column, extra_char):
        current_text = ""
        for index, row in df.iterrows():
            if pd.isna(row[name_column]):
                continue
            cell_value = str(row[name_column]).strip().replace("\t", " ")
            char_value = ""
            if extra_char:
                if name_column + 1 < df.shape[1]:
                    char_value += " " + str(row[name_column + 1]).strip().replace("\t", " ").replace("\n", " ")
                if name_column + 2 < df.shape[1]:
                    char_value += " " + str(row[name_column + 2]).strip().replace("\t", " ").replace("\n", " ")
            else:
                if name_column + 1 < df.shape[1]:
                    char_value = " " + str(row[name_column + 1]).strip().replace("\t", " ").replace("\n", " ")
            if self.is_product_name(cell_value):
                if current_text and self.contains_product_name(current_text):
                    self.data.append({"text": current_text})
                current_text = cell_value + char_value
            else:
                current_text += " " + cell_value + char_value
        if current_text and self.contains_product_name(current_text):
            self.data.append({"text": current_text})

    def process(self):
        self.parse_excel()


class StructuredDocxParser:
    PRODUCT_NAMES = settings.product_names

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []

    def parse_table_type1(self, table):
        """
        Парсинг таблиц по логике первого парсера:
        - Если первая ячейка строки соответствует формату "2.5" – новый товар.
        - Если соответствует формату "2.5.1" – характеристика товара.
        - Иначе – доп. характеристика.
        """
        print("parse_table_type1")
        results = []
        current_product = None
        product_data = {}
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not row_text or not row_text[0]:
                continue
            # Новый товар (пример: 2.5)
            if re.match(r"^\d+\.\d+$", row_text[0]):
                if current_product:
                    results.append(product_data)
                current_product = row_text[0]
                product_data = {"Номенклатура": row_text[0]}
            # Характеристика товара (пример: 2.5.1)
            elif re.match(r"^\d+\.\d+\.\d+$", row_text[0]):
                product_data[row_text[0]] = " | ".join(row_text)
            else:
                product_data[f"Характеристика {len(product_data)}"] = " | ".join(row_text)
        if current_product:
            results.append(product_data)
        return results

    def parse_table_type2(self, table):
        """
        Парсинг таблиц по логике второго парсера:
        - Определяется колонка с заголовком, содержащим "наименование".
        - При нахождении строки с именем товара из PRODUCT_NAMES создаётся новый блок,
          остальные строки добавляются как характеристики.
        """
        print("parse_table_type2")
        results = []
        header_cells = table.rows[0].cells
        name_column = None
        for col_idx, cell in enumerate(header_cells):
            if "наименование" in cell.text.lower():
                name_column = col_idx
                break
        if name_column is None:
            return results

        current_product = None
        product_data = {}
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", " ").replace("\t", " ") for cell in row.cells]
            if name_column < len(row_text):
                product_name = row_text[name_column]
                if any(name.lower() in product_name.lower() for name in self.PRODUCT_NAMES):
                    if current_product:
                        results.append(product_data)
                    current_product = product_name
                    product_data = {"text": product_name, "Характеристики": []}
                elif current_product:
                    # Здесь происходит объединение строки с существующими данными
                    # Для корректного объединения преобразуем список в строку, если нужно
                    additional_text = " ".join(row_text)
                    if isinstance(product_data.get("Характеристики"), list):
                        product_data["Характеристики"].append(additional_text)
                    else:
                        product_data["Характеристики"] = additional_text
        if current_product:
            results.append(product_data)
        return results

    def parse_table_type3(self, table):
        """
        Парсинг таблиц по логике третьего и четвёртого парсеров:
        - Если объединённый текст строки содержит одно из наименований,
          строка считается информационной.
        """
        print("parse_table_type3")
        results = []
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", ";").replace("\t", " ") for cell in row.cells]
            row_combined = " | ".join(row_text)
            if any(name.lower() in row_combined.lower() for name in self.PRODUCT_NAMES):
                # Объединяем данные строки в одну строку
                results.append({"text": row_combined})
        return results

    def parse_paragraphs(self, doc):
        """
        Дополнительная обработка текста вне таблиц (как во втором парсере).
        Из полного текста ищется фрагмент от наименования до слова "гарантия".
        """
        print("parse_paragraphs")
        results = []
        full_text = " ".join([p.text for p in doc.paragraphs])
        lower_text = full_text.lower()
        for name in self.PRODUCT_NAMES:
            if name.lower() in lower_text:
                start_idx = lower_text.find(name.lower())
                end_idx = lower_text.find("гарантия", start_idx)
                product_text = full_text[start_idx:end_idx] if end_idx != -1 else full_text[start_idx:]
                results.append({"text": product_text})
        return results

    def parse_doc(self):
        print(f"Opening file: {self.file_path}")
        if self.file_path.suffix.lower() == ".docx":
            doc = docx.Document(self.file_path)
        elif self.file_path.suffix.lower() == ".doc":
            # Конвертация .doc в .docx с помощью LibreOffice
            temp_docx = self.file_path.with_suffix('.docx')
            try:
                subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'docx', str(self.file_path), '--outdir',
                     str(self.file_path.parent)],
                    check=True
                )
            except subprocess.CalledProcessError as e:
                print("Ошибка при конвертации файла .doc в .docx:", e)
                return
            doc = docx.Document(str(temp_docx))
        else:
            print("Unsupported file format")
            return

        # Обработка таблиц
        for table in doc.tables:
            # Определяем тип таблицы по первой ячейке первого ряда
            if table.rows and table.rows[0].cells:
                first_cell_text = table.rows[0].cells[0].text.lower()
                if "наименование" in first_cell_text:
                    table_type = 2
                elif re.match(r"^\d+\.\d+$", first_cell_text):
                    table_type = 1
                else:
                    table_type = 3
            else:
                continue

            if table_type == 1:
                self.data.extend(self.parse_table_type1(table))
            elif table_type == 2:
                self.data.extend(self.parse_table_type2(table))
            elif table_type == 3:
                self.data.extend(self.parse_table_type3(table))
        # Обработка текста вне таблиц
        self.data.extend(self.parse_paragraphs(doc))

    def _combine_product_data(self, product):
        """
        Объединяет все значения словаря продукта в одну строку.
        Если значение – список, элементы объединяются через пробел.
        """
        parts = []
        for key, value in product.items():
            if isinstance(value, list):
                parts.append(" ".join(value))
            else:
                parts.append(str(value))
        return " ".join(parts)

    def format_output(self):
        """
        Преобразует self.data в список словарей, где каждый словарь имеет единственный ключ "text",
        а значение – объединённая строка со всеми данными по товару.
        """
        formatted_data = []
        for product in self.data:
            combined_text = self._combine_product_data(product)
            formatted_data.append({"text": combined_text})
        self.data = formatted_data
        print(self.data)

    def process(self):
        if not self.file_path.exists():
            print(f"File not found: {self.file_path}")
            return f"File not found: {self.file_path}"
        print(f"Processing file: {self.file_path.name}")
        self.parse_doc()
        self.format_output()


class StructuredPdfParser1:
    PRODUCT_NAMES = settings.product_names

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []

    def parse_pdf(self):
        print(f"Opening file: {self.file_path}")
        with pdfplumber.open(self.file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()
                if not tables:
                    continue

                for table_idx, table in enumerate(tables):
                    print(f"Processing table {table_idx + 1} on page {page_number}")
                    for row in table:
                        row_text = [str(cell).strip().replace("\n", " ") for cell in row if cell]
                        row_combined = " | ".join(row_text)

                        # Проверяем, содержится ли в строке название товара
                        if any(name.lower() in row_combined.lower() for name in self.PRODUCT_NAMES):
                            product_data = {"text": row_combined}
                            self.data.append(product_data)

# TODO: Доделать распознование картинок из pdf
        # if not self.data:
        #     recognition_text = ''
        #     try:
        #         pages = convert_from_path(str(self.file_path))
        #         for page in pages:
        #             recognition_text += pytesseract.image_to_string(page, lang='rus') + "\n"
        #     except Exception as e:
        #         print("OCR error:", e)
        #     print(recognition_text)
        #     self.data.append(recognition_text.strip())

    def print_data(self):
        if not self.data:
            print("No data parsed!")
        for product in self.data:
            print(product)

    def process(self):
        if not self.file_path.exists():
            print(f"File not found: {self.file_path}")
            return

        print(f"Processing file: {self.file_path.name}")
        self.parse_pdf()


class StructuredPdfParser2:
    PRODUCT_NAMES = settings.product_names
    EXCLUDE_WORDS = ["шт.", "шт", "штук"]

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []
        self.header_candidates = []  # список токенов из первых пяти найденных заголовков товара
        self.header_mask = None      # регулярное выражение для определения начала нового товара

    # def determine_common_pattern(self):
    #     if not self.header_candidates:
    #         return None
    #
    #     # Проверяем, сколько кандидатов соответствуют схеме "число + точка"
    #     digit_dot_pattern = re.compile(r'^\d+\.$')
    #     count_digit_dot = sum(1 for token in self.header_candidates if digit_dot_pattern.match(token))
    #     if count_digit_dot >= 3:
    #         return re.compile(r'^\s*\d+\s*\.')
    #
    #     # Иначе выбираем наиболее часто встречающийся токен
    #     counter = collections.Counter(self.header_candidates)
    #     most_common_token, freq = counter.most_common(1)[0]
    #     if most_common_token:
    #         escaped = re.escape(most_common_token)
    #         return re.compile(r'^\s*' + escaped)
    #     return None
    #
    # def update_header_mask(self):
    #     if len(self.header_candidates) >= 5 and not self.header_mask:
    #         self.header_mask = self.determine_common_pattern()

    # def is_new_header(self, row_combined):
    #     """
    #     Если паттерн определён, строка считается началом нового товара, если совпадает с паттерном.
    #     Если паттерн не установлен, в качестве критерия используется наличие наименования товара.
    #     """
    #     if self.header_mask:
    #         return bool(self.header_mask.match(row_combined))
    #     return any(name.lower() in row_combined.lower() for name in self.PRODUCT_NAMES)

    def parse_pdf(self):
        print(f"Opening file: {self.file_path}")
        current_record = None
        with pdfplumber.open(self.file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()
                if not tables:
                    continue

                for table_idx, table in enumerate(tables):
                    for row in table:
                        row_text = [str(cell).strip().replace("\n", " ") for cell in row if cell]
                        if not row_text:
                            continue
                        row_combined = " | ".join(row_text)
                        row_lower = row_combined.lower()

                        # Пропускаем строки, содержащие исключающие слова
                        if any(ex_word in row_lower for ex_word in self.EXCLUDE_WORDS):
                            continue

                        # Если строка содержит наименование товара, считаем её кандидатом на начало записи
                        if any(name.lower() in row_lower for name in self.PRODUCT_NAMES):
                            tokens = row_combined.split()
                            first_token = tokens[0] if tokens else ""
                            if first_token and first_token not in self.header_candidates:
                                self.header_candidates.append(first_token)
                            # Обновляем маску, если набрано достаточно кандидатов
                            self.update_header_mask()

                            # Если уже есть накопленная запись, сохраняем её
                            if current_record:
                                self.data.append({"text": current_record})
                            current_record = row_combined
                        else:
                            # Если паттерн установлен и строка соответствует началу нового товара, то новый заголовок
                            if self.header_mask and self.is_new_header(row_combined):
                                if current_record:
                                    self.data.append({"text": current_record})
                                current_record = row_combined
                            else:
                                # Иначе строка считается продолжением предыдущего товара
                                if current_record:
                                    if current_record.endswith('-'):
                                        current_record = current_record.rstrip('-') + row_combined.lstrip()
                                    else:
                                        current_record += " " + row_combined
        if current_record:
            self.data.append({"text": current_record})

    def process(self):
        if not self.file_path.exists():
            print(f"File not found: {self.file_path}")
            return

        print(f"Processing file: {self.file_path.name}")
        self.parse_pdf()


class StructuredPdfParser3:
    PRODUCT_NAMES = settings.product_names
    EXCLUDE_WORDS = ["шт.", "шт", "штук"]

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []
        self.header_candidates = []  # список токенов из первых найденных заголовков товара
        self.header_mask = None  # регулярное выражение для определения начала нового товара

    def determine_common_pattern(self):
        if not self.header_candidates:
            return None

        # Проверяем, сколько кандидатов соответствуют схеме "число + точка"
        digit_dot_pattern = re.compile(r'^\d+\.$')
        count_digit_dot = sum(1 for token in self.header_candidates if digit_dot_pattern.match(token))
        if count_digit_dot >= 3:
            return re.compile(r'^\s*\d+\s*\.')

        # Иначе выбираем наиболее часто встречающийся токен
        counter = collections.Counter(self.header_candidates)
        most_common_token, freq = counter.most_common(1)[0]
        if most_common_token:
            escaped = re.escape(most_common_token)
            return re.compile(r'^\s*' + escaped)
        return None

    def update_header_mask(self):
        if len(self.header_candidates) >= 5 and not self.header_mask:
            self.header_mask = self.determine_common_pattern()

    def is_new_header(self, row_combined):
        """
        Если паттерн определён, строка считается началом нового товара, если совпадает с паттерном.
        Если паттерн не установлен, в качестве критерия используется наличие наименования товара.
        """
        if self.header_mask:
            return bool(self.header_mask.match(row_combined))
        return any(name.lower() in row_combined.lower() for name in self.PRODUCT_NAMES)

    def parse_pdf(self):
        print(f"Открытие файла: {self.file_path}")
        current_record = None
        try:
            # Конвертируем страницы PDF в изображения
            images = convert_from_path(self.file_path)
        except Exception as e:
            print(f"Ошибка при конвертации PDF в изображения: {e}")
            return

        for page_number, image in enumerate(images, start=1):
            print(f"Обработка страницы {page_number} с помощью OCR")
            # Извлекаем текст с изображения; при необходимости можно указать язык (например, lang='rus')
            page_text = pytesseract.image_to_string(image, lang='rus')
            lines = page_text.splitlines()
            for line in lines:
                row_combined = line.strip()
                if not row_combined:
                    continue
                row_lower = row_combined.lower()

                # Пропускаем строки, содержащие исключающие слова
                if any(ex_word in row_lower for ex_word in self.EXCLUDE_WORDS):
                    continue

                # Если строка содержит наименование товара, считаем её кандидатом на начало записи
                if any(name.lower() in row_lower for name in self.PRODUCT_NAMES):
                    tokens = row_combined.split()
                    first_token = tokens[0] if tokens else ""
                    if first_token and first_token not in self.header_candidates:
                        self.header_candidates.append(first_token)
                    self.update_header_mask()

                    # Если уже есть накопленная запись, сохраняем её
                    if current_record:
                        self.data.append({"text": current_record})
                    current_record = row_combined
                else:
                    # Если паттерн установлен и строка соответствует началу нового товара, то считаем её новым заголовком
                    if self.header_mask and self.is_new_header(row_combined):
                        if current_record:
                            self.data.append({"text": current_record})
                        current_record = row_combined
                    else:
                        # Иначе строка считается продолжением предыдущего товара
                        if current_record:
                            if current_record.endswith('-'):
                                current_record = current_record.rstrip('-') + row_combined.lstrip()
                            else:
                                current_record += " " + row_combined

        if current_record:
            self.data.append({"text": current_record})

    def process(self):
        if not self.file_path.exists():
            print(f"Файл не найден: {self.file_path}")
            return

        print(f"Обработка файла: {self.file_path.name}")
        self.parse_pdf()
