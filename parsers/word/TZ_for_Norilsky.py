import docx
from pathlib import Path
import importlib.util

BASE_DIR = Path(__file__).resolve().parents[2]
SETTINGS_PATH = BASE_DIR / "settings.py"

spec = importlib.util.spec_from_file_location("settings", SETTINGS_PATH)
settings = importlib.util.module_from_spec(spec)
spec.loader.exec_module(settings)

# Проблема совместимости данных!!! Нужно будет подключать модель параллельно с парсингом для проверки данных


class StructuredDocxParser:
    PRODUCT_NAMES = settings.product_names

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []

    def parse_doc(self):
        print(f"Opening file: {self.file_path}")
        doc = docx.Document(self.file_path)

        for table in doc.tables:
            name_column = None
            for col_idx in range(len(table.rows[0].cells)):
                if any("наименование" in table.rows[0].cells[col_idx].text.lower() for _ in table.rows):
                    name_column = col_idx
                    break

            if name_column is None:
                continue

            current_product = None
            product_data = {}

            for row in table.rows:
                row_text = [cell.text.strip().replace("\n", " ").replace("\t", " ") for cell in row.cells]
                row_combined = " | ".join(row_text)
                print(f"Row text: {row_combined}")

                if name_column < len(row_text):
                    product_name = row_text[name_column]
                    if any(name.lower() in product_name.lower() for name in self.PRODUCT_NAMES):
                        if current_product:
                            self.data.append(product_data)
                        current_product = product_name
                        product_data = {"0": current_product, "Характеристики": []}
                    elif current_product:
                        product_data["Характеристики"].append(" ".join(row_text))

            if current_product:
                self.data.append(product_data)

        # Обработка текста вне таблиц
        full_text = " ".join([p.text for p in doc.paragraphs])
        for name in self.PRODUCT_NAMES:
            if name.lower() in full_text.lower():
                start_idx = full_text.lower().find(name.lower())
                end_idx = full_text.lower().find("гарантия", start_idx)
                product_text = full_text[start_idx:end_idx] if end_idx != -1 else full_text[start_idx:]
                self.data.append({"0": name, "Характеристики": product_text})

    def print_data(self):
        if not self.data:
            print("No data parsed!")
        for product in self.data:
            print(product)

    def process(self):
        if not self.file_path.exists():
            print(f"File not found: {self.file_path}")
            return

        print(f"Processing file: {self.file_path.name}")
        self.parse_doc()
        self.print_data()


# Использование
file_path = Path("..", "..", "test_data", "input", "ТЗ Норильский транспорт.docx")  # Конкретный файл DOC
parser = StructuredDocxParser(file_path)
parser.process()

# Проблема совместимости данных!!! Нужно будет подключать модель параллельно с парсингом для проверки данных