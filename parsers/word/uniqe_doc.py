import docx
import re
from pathlib import Path
import importlib.util

# Общие настройки и пути
BASE_DIR = Path(__file__).resolve().parents[2]
SETTINGS_PATH = BASE_DIR / "settings.py"

spec = importlib.util.spec_from_file_location("settings", SETTINGS_PATH)
settings = importlib.util.module_from_spec(spec)
spec.loader.exec_module(settings)


class StructuredDocxParser:
    PRODUCT_NAMES = settings.product_names

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.data = []

    def parse_table_type1(self, table):
        """
        Парсинг таблиц по логике первого парсера:
        - Если первая ячейка строки соответствует формату "2.5" – новый товар.
        - Если соответствует формату "2.5.1" – характеристика товара.
        - Иначе – доп. характеристика.
        """
        results = []
        current_product = None
        product_data = {}
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not row_text or not row_text[0]:
                continue
            # Новый товар (пример: 2.5)
            if re.match(r"^\d+\.\d+$", row_text[0]):
                if current_product:
                    results.append(product_data)
                current_product = row_text[0]
                product_data = {"Номенклатура": row_text[0]}
            # Характеристика товара (пример: 2.5.1)
            elif re.match(r"^\d+\.\d+\.\d+$", row_text[0]):
                product_data[row_text[0]] = " | ".join(row_text)
            else:
                product_data[f"Характеристика {len(product_data)}"] = " | ".join(row_text)
        if current_product:
            results.append(product_data)
        return results

    def parse_table_type2(self, table):
        """
        Парсинг таблиц по логике второго парсера:
        - Определяется колонка с заголовком, содержащим "наименование".
        - При нахождении строки с именем товара из PRODUCT_NAMES создаётся новый блок,
          остальные строки добавляются как характеристики.
        """
        results = []
        header_cells = table.rows[0].cells
        name_column = None
        for col_idx, cell in enumerate(header_cells):
            if "наименование" in cell.text.lower():
                name_column = col_idx
                break
        if name_column is None:
            return results

        current_product = None
        product_data = {}
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", " ").replace("\t", " ") for cell in row.cells]
            if name_column < len(row_text):
                product_name = row_text[name_column]
                if any(name.lower() in product_name.lower() for name in self.PRODUCT_NAMES):
                    if current_product:
                        results.append(product_data)
                    current_product = product_name
                    product_data = {"0": product_name, "Характеристики": []}
                elif current_product:
                    product_data["Характеристики"].append(" ".join(row_text))
        if current_product:
            results.append(product_data)
        return results

    def parse_table_type3(self, table):
        """
        Парсинг таблиц по логике третьего и четвёртого парсеров:
        - Если объединённый текст строки содержит одно из наименований,
          строка считается информационной.
        """
        results = []
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", ";").replace("\t", " ") for cell in row.cells]
            row_combined = " | ".join(row_text)
            if any(name.lower() in row_combined.lower() for name in self.PRODUCT_NAMES):
                product_data = {"0": row_text}
                results.append(product_data)
        return results

    def parse_paragraphs(self, doc):
        """
        Дополнительная обработка текста вне таблиц (как во втором парсере).
        Из полного текста ищется фрагмент от наименования до слова "гарантия".
        """
        results = []
        full_text = " ".join([p.text for p in doc.paragraphs])
        lower_text = full_text.lower()
        for name in self.PRODUCT_NAMES:
            if name.lower() in lower_text:
                start_idx = lower_text.find(name.lower())
                end_idx = lower_text.find("гарантия", start_idx)
                product_text = full_text[start_idx:end_idx] if end_idx != -1 else full_text[start_idx:]
                results.append({"0": name, "Характеристики": product_text})
        return results

    def parse_doc(self):
        print(f"Opening file: {self.file_path}")
        doc = docx.Document(self.file_path)
        # Обработка таблиц
        for table in doc.tables:
            # Определяем тип таблицы по первой ячейке первого ряда
            if table.rows and table.rows[0].cells:
                first_cell_text = table.rows[0].cells[0].text.lower()
                if "наименование" in first_cell_text:
                    table_type = 2
                elif re.match(r"^\d+\.\d+$", first_cell_text):
                    table_type = 1
                else:
                    table_type = 3
            else:
                continue

            if table_type == 1:
                self.data.extend(self.parse_table_type1(table))
            elif table_type == 2:
                self.data.extend(self.parse_table_type2(table))
            elif table_type == 3:
                self.data.extend(self.parse_table_type3(table))
        # Обработка текста вне таблиц
        self.data.extend(self.parse_paragraphs(doc))

    def print_data(self):
        if not self.data:
            print("No data parsed!")
        for product in self.data:
            print(product)

    def process(self):
        print()
        if not self.file_path.exists():
            print(f"File not found: {self.file_path}")
            return
        print(f"Processing file: {self.file_path.name}")
        self.parse_doc()
        self.print_data()


# 1
file_path = Path("..", "..", "test_data", "input", "ТЗ для МГУ.docx")
parser = StructuredDocxParser(file_path)
parser.process()

# 2
file_path = Path("..", "..", "test_data", "input", "ТЗ для Рос Волга.docx")
parser = StructuredDocxParser(file_path)
parser.process()

# 3
file_path = Path("..", "..", "test_data", "input", "ТЗ для Туапсе.docx")
parser = StructuredDocxParser(file_path)
parser.process()

# 4
file_path = Path("..", "..", "test_data", "input", "ТЗ Норильский транспорт.docx")
parser = StructuredDocxParser(file_path)
parser.process()
